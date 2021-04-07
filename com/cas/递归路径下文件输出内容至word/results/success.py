import os
from docx import Document
# import sys
# import os
# 1、os.listdir('/root/python/') #列出当前目录下所有文件
# 2、os.path.isdir('/abc') #判断是否是目录，返回布尔值，不存在也返回false
# 3、os.path.isfile('/etc/passwd') #判断是否是文件
# 4、os.path.join('/etc/', 'passwd') #连接文件，返回/etc/passwd
# 参数准备

document = Document()
document.add_heading('源代码', 0)
basedir = '/Users/xianglong/IdeaProjects/rabbitmq-service-backend'

def print_files(path):
    global document
    # print(basedir)
    # 第一步过滤，拿到所有的文件夹
    lsdir = os.listdir(path)

    dirs = [i for i in lsdir if os.path.isdir(os.path.join(
        path, i))]
    if dirs:
        for i in dirs:
            print_files(os.path.join(path, i))
    files = [i for i in lsdir if os.path.isfile(os.path.join(path,i)) and os.path.join(path,i).endswith("java")]
    for f in files:
        print (os.path.join(path, f))
        fo = open(os.path.join(path, f), "r+")
        document.add_paragraph(fo.read())
        fo.close()
    document.save('源代码.docx')

# print_files(sys.argv[1])


if __name__ == '__main__':
    print("文件递归开始....")
    print_files(basedir)
